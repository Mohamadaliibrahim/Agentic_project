"""
Document Processing Service
Handles PDF, CSV, Word, and text file processing into chunks using tiktoken cl100k_base encoding
Saves processed documents as .pkl files in vector_storage folder (same as rag_testing notebook)
"""

import os
import uuid
import pickle
import logging
from typing import List, Dict, Any
from io import BytesIO
import tempfile
import datetime
import numpy as np

try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import csv
    from io import StringIO
except ImportError:
    csv = None
    StringIO = None

try:
    import tiktoken
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    tiktoken = None
    RecursiveCharacterTextSplitter = None

try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    faiss = None
    FAISS_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing various document types into text chunks using tiktoken and saving as .pkl files"""
    
    def __init__(self, chunk_size: int = 600, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.encoding_name = 'cl100k_base'
        self.separators = ["###", "\n\n\n", "\n\n", "\n"]
        
        if tiktoken:
            self.encoder = tiktoken.get_encoding(self.encoding_name)
        else:
            logger.warning("tiktoken not available, falling back to character-based chunking")
            self.encoder = None
        
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.csv': self._process_csv,
            '.docx': self._process_docx,
            '.doc': self._process_docx
        }
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        user_id: str
    ) -> Dict[str, Any]:
        """
        Process document and return chunks with metadata
        
        Args:
            file_content: The file content as bytes
            filename: Original filename
            user_id: User who uploaded the document
            
        Returns:
            Dict containing document info and text chunks
        """
        try:
            doc_id = str(uuid.uuid4())
            
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            processor = self.supported_types[file_ext]
            text_content = await processor(file_content, filename)
            
            if text_content is None:
                raise Exception(f"Text extraction returned None for {filename}")
            
            if not isinstance(text_content, str):
                raise Exception(f"Text extraction returned invalid type for {filename}: {type(text_content)}")
            
            if not text_content.strip():
                raise Exception(f"No text content extracted from {filename}")
            
            chunks = self._create_chunks(text_content)
            
            if not chunks:
                raise Exception(f"No chunks created from {filename}")
            
            document_info = {
                "document_id": doc_id,
                "filename": filename,
                "file_type": file_ext,
                "user_id": user_id,
                "total_chunks": len(chunks),
                "total_characters": len(text_content),
                "chunks": chunks
            }
            
            logger.info(f"Processed document {filename}: {len(chunks)} chunks created")
            return document_info
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> str:
        """Process PDF file and extract text"""
        if not PyPDF2:
            raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF file has no pages")
            
            text_content = ""
            pages_with_text = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text
                        pages_with_text += 1
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1} in {filename}: {e}")
                    continue
            
            if pages_with_text == 0:
                raise Exception(f"No readable text found in PDF '{filename}'. The PDF might be image-based or corrupted.")
            
            text_content = text_content.strip()
            if not text_content:
                raise Exception(f"PDF '{filename}' contains only empty pages")
                
            logger.info(f"Extracted text from {pages_with_text} pages of {filename}")
            return text_content
            
        except Exception as e:
            if "No readable text found" in str(e) or "contains only empty pages" in str(e):
                raise e 
            else:
                raise Exception(f"Error processing PDF '{filename}': {str(e)}")
    
    async def _process_text(self, file_content: bytes, filename: str) -> str:
        """Process plain text file"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    return text_content
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any common encoding")
            
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    async def _process_csv(self, file_content: bytes, filename: str) -> str:
        """Process CSV file and convert to readable text"""
        if not pd:
            raise Exception("pandas not installed. Install with: pip install pandas")
        
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    csv_text = file_content.decode(encoding)
                    csv_file = StringIO(csv_text)
                    df = pd.read_csv(csv_file)
                    
                    text_content = f"CSV File: {filename}\n"
                    text_content += f"Columns: {', '.join(df.columns.tolist())}\n"
                    text_content += f"Number of rows: {len(df)}\n\n"
                    
                    text_content += "Column Information:\n"
                    for col in df.columns:
                        text_content += f"- {col}: {df[col].dtype}\n"
                    
                    text_content += "\nData Summary:\n"
                    text_content += df.describe(include='all').to_string()
                    

                    text_content += f"\n\nFirst 10 rows:\n"
                    text_content += df.head(10).to_string()
                    
                    return text_content
                    
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    if encoding == encodings[-1]:
                        raise e
                    continue
            
            raise Exception("Could not decode CSV file with any common encoding")
            
        except Exception as e:
            raise Exception(f"Error processing CSV: {str(e)}")
    
    async def _process_docx(self, file_content: bytes, filename: str) -> str:
        """Process Word document (DOCX)"""
        if not DocxDocument:
            raise Exception("python-docx not installed. Install with: pip install python-docx")
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"
                
                for table in doc.tables:
                    text_content += "\n--- Table ---\n"
                    for row in table.rows:
                        row_text = "\t".join([cell.text for cell in row.cells])
                        text_content += row_text + "\n"
                
                if not text_content.strip():
                    raise Exception("No text could be extracted from Word document")
                
                return text_content
                
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks using tiktoken cl100k_base encoding"""
        chunks = []
        
        try:
            if RecursiveCharacterTextSplitter and self.encoder:
                text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                    chunk_size=self.chunk_size, 
                    chunk_overlap=self.overlap, 
                    separators=self.separators,
                    encoding_name=self.encoding_name
                )
                
                chunk_texts = text_splitter.split_text(text)
                
                for i, chunk_text in enumerate(chunk_texts):
                    if chunk_text.strip():
                        token_count = len(self.encoder.encode(chunk_text))
                        
                        chunk_info = {
                            "chunk_id": str(uuid.uuid4()),
                            "text": chunk_text,
                            "chunk_index": i,
                            "word_count": len(chunk_text.split()),
                            "character_count": len(chunk_text),
                            "token_count": token_count
                        }
                        chunks.append(chunk_info)
                
                logger.info(f"Created {len(chunks)} chunks using tiktoken cl100k_base encoding")
                
            else:
                logger.warning("tiktoken or langchain not available, using simple word-based chunking")
                words = text.split()
                
                for i in range(0, len(words), self.chunk_size - self.overlap):
                    chunk_words = words[i:i + self.chunk_size]
                    chunk_text = ' '.join(chunk_words)
                    
                    if chunk_text.strip():
                        chunk_info = {
                            "chunk_id": str(uuid.uuid4()),
                            "text": chunk_text,
                            "chunk_index": len(chunks),
                            "word_count": len(chunk_words),
                            "character_count": len(chunk_text),
                            "token_count": len(chunk_words)
                        }
                        chunks.append(chunk_info)
                
        except Exception as e:
            logger.error(f"Error in tiktoken chunking, falling back to simple chunking: {str(e)}")
            words = text.split()
            
            for i in range(0, len(words), self.chunk_size - self.overlap):
                chunk_words = words[i:i + self.chunk_size]
                chunk_text = ' '.join(chunk_words)
                
                if chunk_text.strip():
                    chunk_info = {
                        "chunk_id": str(uuid.uuid4()),
                        "text": chunk_text,
                        "chunk_index": len(chunks),
                        "word_count": len(chunk_words),
                        "character_count": len(chunk_text),
                        "token_count": len(chunk_words)
                    }
                    chunks.append(chunk_info)
        
        return chunks
    
    async def process_and_store_document(
        self,
        file_content: bytes,
        filename: str,
        user_id: str
    ) -> Dict[str, Any]:
        """
        Process document and store in shared user vector_storage file like rag_testing notebook
        All user documents go into one shared .pkl file for the user
        
        Args:
            file_content: The file content as bytes
            filename: Original filename
            user_id: User who uploaded the document
            
        Returns:
            Dict containing document info and storage details
        """
        try:
            document_info = await self.process_document(file_content, filename, user_id)
            
            chunks = document_info["chunks"]
            chunk_texts = [chunk["text"] for chunk in chunks]
            
            from core.embedding_service import embedding_service
            embeddings = await embedding_service.generate_embeddings(chunk_texts)
            
            embeddings_array = np.array(embeddings, dtype=np.float32)
            norms = np.linalg.norm(embeddings_array, axis=1, keepdims=True)
            norms[norms == 0] = 1e-9
            normalized_embeddings = embeddings_array / norms
            
            vector_dir = f"vector_storage/{user_id}"
            os.makedirs(vector_dir, exist_ok=True)
            
            index_file = os.path.join(vector_dir, "index.faiss")
            metadata_file = os.path.join(vector_dir, "metadata.pkl")
            
            existing_metadata = []
            existing_embeddings = None
            
            if os.path.exists(metadata_file) and os.path.exists(index_file):
                try:
                    with open(metadata_file, "rb") as f:
                        existing_metadata = pickle.load(f)
                    
                    existing_index = faiss.read_index(index_file)
                    
                    if existing_index.ntotal > 0:
                        existing_embeddings = np.zeros((existing_index.ntotal, existing_index.d), dtype=np.float32)
                        
                        try:
                            existing_index.reconstruct_n(0, existing_index.ntotal, existing_embeddings)
                            logger.info(f"Loaded existing {len(existing_metadata)} chunks for user {user_id}")
                        except Exception as reconstruct_error:
                            logger.warning(f"Could not reconstruct embeddings from existing index: {reconstruct_error}. Will rebuild index.")
                            if existing_metadata and all('embedding' in meta for meta in existing_metadata):
                                existing_embeddings = np.array([meta['embedding'] for meta in existing_metadata], dtype=np.float32)
                                logger.info(f"Recovered embeddings from metadata for {len(existing_metadata)} chunks")
                            else:
                                logger.warning("No embeddings available in metadata. Creating new index.")
                                existing_metadata = []
                                existing_embeddings = None
                    else:
                        logger.info("Existing index is empty. Starting fresh.")
                        existing_metadata = []
                        existing_embeddings = None
                    
                except Exception as e:
                    logger.warning(f"Could not load existing data: {e}. Creating new index.")
                    existing_metadata = []
                    existing_embeddings = None
            
            published_date = datetime.datetime.utcnow().isoformat()
            doc_id = document_info["document_id"]
            
            new_metadata = []
            for i, chunk in enumerate(chunks):
                new_metadata.append({
                    "content": chunk["text"],
                    "embedding": normalized_embeddings[i].tolist(),
                    "published_date": published_date,
                    "filename": filename,
                    "document_id": doc_id,
                    "user_id": user_id,
                    "chunk_id": chunk["chunk_id"],
                    "chunk_index": i,
                    "url": f"document://{doc_id}"
                })
            
            all_metadata = existing_metadata + new_metadata
            
            if existing_embeddings is not None:
                all_embeddings = np.vstack([existing_embeddings, normalized_embeddings])
            else:
                all_embeddings = normalized_embeddings
            
            all_embeddings = all_embeddings.astype(np.float32)
            
            if not FAISS_AVAILABLE:
                raise Exception("FAISS not available. Install with: pip install faiss-cpu")
            
            dim = all_embeddings.shape[1]
            index = faiss.IndexFlatIP(dim)
            index.add(all_embeddings)
            
            faiss.write_index(index, index_file)
            with open(metadata_file, "wb") as f:
                pickle.dump(all_metadata, f)
            
            logger.info(f"Updated shared storage: {len(new_metadata)} new chunks added. Total: {len(all_metadata)} chunks")
            
            from database.factory import get_db
            db = get_db()
            
            doc_data = {
                "document_id": doc_id,
                "filename": filename,
                "file_type": document_info["file_type"],
                "user_id": user_id,
                "upload_date": datetime.datetime.utcnow(),
                "total_chunks": len(chunks),
                "vector_storage_path": vector_dir,
                "index_file": index_file,
                "metadata_file": metadata_file
            }
            await db.store_document(doc_data)
            
            chunk_data = []
            start_idx = len(existing_metadata)
            for i, chunk in enumerate(chunks):
                chunk_doc = {
                    "chunk_id": chunk["chunk_id"],
                    "document_id": doc_id,
                    "user_id": user_id,
                    "text": chunk["text"],
                    "embedding": normalized_embeddings[i].tolist(),
                    "chunk_index": start_idx + i,
                    "metadata": new_metadata[i]
                }
                chunk_data.append(chunk_doc)
            
            await db.store_document_chunks(chunk_data)
            
            return {
                "document_id": doc_id,
                "filename": filename,
                "file_type": document_info["file_type"],
                "total_chunks": len(chunks),
                "vector_storage_path": vector_dir,
                "index_file": index_file,
                "metadata_file": metadata_file,
                "total_user_chunks": len(all_metadata),
                "message": f"Document processed and added to shared storage. {len(chunks)} new chunks added. Total user chunks: {len(all_metadata)}"
            }
            
        except Exception as e:
            logger.error(f"Error processing and storing document {filename}: {str(e)}")
            raise Exception(f"Failed to process and store document: {str(e)}")

    async def search_user_documents(self, user_id: str, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        Search user's documents using shared FAISS index like in rag_testing notebook
        All user documents are in one shared index file
        
        Args:
            user_id: User identifier
            query: Search query
            top_k: Number of results to return
            
        Returns:
            List of search results with scores and metadata
        """
        try:
            vector_dir = f"vector_storage/{user_id}"
            if not os.path.exists(vector_dir):
                logger.warning(f"No vector storage found for user {user_id}")
                return []
            
            index_file = os.path.join(vector_dir, "index.faiss")
            metadata_file = os.path.join(vector_dir, "metadata.pkl")
            
            if not os.path.exists(index_file) or not os.path.exists(metadata_file):
                logger.warning(f"No shared index files found for user {user_id}")
                return []
            
            from core.embedding_service import embedding_service
            query_embeddings = await embedding_service.generate_embeddings([query])
            query_embedding = np.array(query_embeddings[0])
            
            query_norm = query_embedding / (np.linalg.norm(query_embedding) + 1e-9)
            query_norm = query_norm.reshape(1, -1)
            
            try:
                index = faiss.read_index(index_file)
                with open(metadata_file, "rb") as f:
                    metadata = pickle.load(f)
                
                D, I = index.search(query_norm, min(top_k, len(metadata)))
                
                results = []
                
                for distance, idx in zip(D[0], I[0]):
                    if idx >= 0 and idx < len(metadata):
                        similarity = float(distance)
                        results.append({
                            "score": similarity,
                            "metadata": metadata[idx],
                            "document_id": metadata[idx]["document_id"]
                        })
                
                results.sort(key=lambda x: x["score"], reverse=True)
                logger.info(f"Found {len(results)} search results for query '{query}' for user {user_id}")
                return results
                
            except Exception as e:
                logger.error(f"Error searching shared index for user {user_id}: {str(e)}")
                return []
            
        except Exception as e:
            logger.error(f"Error searching documents for user {user_id}: {str(e)}")
            return []

document_processor = DocumentProcessor()
