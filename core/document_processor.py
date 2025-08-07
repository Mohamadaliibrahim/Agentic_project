"""
Document Processing Service
Handles PDF, CSV, Word, and text file processing into chunks
"""

import os
import uuid
import logging
from typing import List, Dict, Any
from io import BytesIO
import tempfile

try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import csv
    from io import StringIO
except ImportError:
    csv = None
    StringIO = None

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing various document types into text chunks"""
    
    def __init__(self, chunk_size: int = 600, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.csv': self._process_csv,
            '.docx': self._process_docx,
            '.doc': self._process_docx
        }
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        user_id: str
    ) -> Dict[str, Any]:
        """
        Process document and return chunks with metadata
        
        Args:
            file_content: The file content as bytes
            filename: Original filename
            user_id: User who uploaded the document
            
        Returns:
            Dict containing document info and text chunks
        """
        try:
            # Generate document ID
            doc_id = str(uuid.uuid4())
            
            # Get file extension
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Process document based on type
            processor = self.supported_types[file_ext]
            text_content = await processor(file_content, filename)
            
            # Split into chunks
            chunks = self._create_chunks(text_content)
            
            # Create document metadata
            document_info = {
                "document_id": doc_id,
                "filename": filename,
                "file_type": file_ext,
                "user_id": user_id,
                "total_chunks": len(chunks),
                "total_characters": len(text_content),
                "chunks": chunks
            }
            
            logger.info(f"Processed document {filename}: {len(chunks)} chunks created")
            return document_info
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> str:
        """Process PDF file and extract text"""
        if not PyPDF2:
            raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content.strip():
                raise Exception("No text could be extracted from PDF")
                
            return text_content
            
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    async def _process_text(self, file_content: bytes, filename: str) -> str:
        """Process plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    return text_content
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any common encoding")
            
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    async def _process_csv(self, file_content: bytes, filename: str) -> str:
        """Process CSV file and convert to readable text"""
        if not pd:
            raise Exception("pandas not installed. Install with: pip install pandas")
        
        try:
            # Try different encodings for CSV
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    csv_text = file_content.decode(encoding)
                    csv_file = StringIO(csv_text)
                    df = pd.read_csv(csv_file)
                    
                    # Convert DataFrame to readable text
                    text_content = f"CSV File: {filename}\n"
                    text_content += f"Columns: {', '.join(df.columns.tolist())}\n"
                    text_content += f"Number of rows: {len(df)}\n\n"
                    
                    # Add column descriptions
                    text_content += "Column Information:\n"
                    for col in df.columns:
                        text_content += f"- {col}: {df[col].dtype}\n"
                    
                    text_content += "\nData Summary:\n"
                    text_content += df.describe(include='all').to_string()
                    
                    # Add first few rows as examples
                    text_content += f"\n\nFirst 10 rows:\n"
                    text_content += df.head(10).to_string()
                    
                    return text_content
                    
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    if encoding == encodings[-1]:  # Last encoding attempt
                        raise e
                    continue
            
            raise Exception("Could not decode CSV file with any common encoding")
            
        except Exception as e:
            raise Exception(f"Error processing CSV: {str(e)}")
    
    async def _process_docx(self, file_content: bytes, filename: str) -> str:
        """Process Word document (DOCX)"""
        if not DocxDocument:
            raise Exception("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Save to temporary file since python-docx needs file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"
                
                # Extract table content
                for table in doc.tables:
                    text_content += "\n--- Table ---\n"
                    for row in table.rows:
                        row_text = "\t".join([cell.text for cell in row.cells])
                        text_content += row_text + "\n"
                
                if not text_content.strip():
                    raise Exception("No text could be extracted from Word document")
                
                return text_content
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        chunks = []
        
        # Simple chunking strategy
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if chunk_text.strip():
                chunk_info = {
                    "chunk_id": str(uuid.uuid4()),
                    "text": chunk_text,
                    "chunk_index": len(chunks),
                    "word_count": len(chunk_words),
                    "character_count": len(chunk_text)
                }
                chunks.append(chunk_info)
        
        return chunks

# Create singleton instance
document_processor = DocumentProcessor()
